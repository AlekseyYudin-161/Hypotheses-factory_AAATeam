from __future__ import annotations

import hashlib
import json
import mimetypes
import re
from pathlib import Path
from typing import Any, Iterable

import pandas as pd
from docx import Document
from pypdf import PdfReader

MAX_FILES = 20
MAX_FILE_BYTES = 20 * 1024 * 1024
MAX_TOTAL_BYTES = 60 * 1024 * 1024
MAX_TEXT_CHARS = 30_000
MAX_EXCERPT_CHARS = 700


def _safe_name(name: str) -> str:
    original = Path(name)
    stem = re.sub(r"[^A-Za-zА-Яа-яЁё0-9._-]+", "_", original.stem).strip("._") or "document"
    suffix = original.suffix.lower()
    return f"{stem[:100]}{suffix}"


def _clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _extract_pdf(path: Path) -> tuple[str, str, int | None]:
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
        if sum(len(part) for part in parts) >= MAX_TEXT_CHARS:
            break
    text = "\n".join(parts).strip()
    status = "text_extracted" if text else "requires_ocr"
    return text[:MAX_TEXT_CHARS], status, len(reader.pages)


def _extract_docx(path: Path) -> tuple[str, str, int | None]:
    document = Document(str(path))
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
    text = "\n".join(blocks)
    return text[:MAX_TEXT_CHARS], "text_extracted", None


def _extract_text_file(path: Path) -> tuple[str, str, int | None]:
    raw = path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            text = raw.decode(encoding)
            return text[:MAX_TEXT_CHARS], "text_extracted", None
        except UnicodeDecodeError:
            continue
    return "", "read_error", None


def _extract_csv(path: Path) -> tuple[str, str, int | None]:
    dataframe = None
    for encoding in ("utf-8", "utf-8-sig", "cp1251"):
        try:
            dataframe = pd.read_csv(path, encoding=encoding, sep=None, engine="python")
            break
        except Exception:
            continue
    if dataframe is None:
        return "", "read_error", None
    return dataframe.head(500).to_csv(index=False)[:MAX_TEXT_CHARS], "table_extracted", None


def _extract_xlsx(path: Path) -> tuple[str, str, int | None]:
    sheets = pd.read_excel(path, sheet_name=None)
    parts: list[str] = []
    for sheet_name, dataframe in sheets.items():
        parts.append(f"SHEET: {sheet_name}\n{dataframe.head(300).to_csv(index=False)}")
        if sum(len(part) for part in parts) >= MAX_TEXT_CHARS:
            break
    return "\n\n".join(parts)[:MAX_TEXT_CHARS], "table_extracted", None


def _extract_json(path: Path) -> tuple[str, str, int | None]:
    try:
        data = json.loads(path.read_text(encoding="utf-8-sig"))
        text = json.dumps(data, ensure_ascii=False, indent=2)
        return text[:MAX_TEXT_CHARS], "text_extracted", None
    except Exception:
        return _extract_text_file(path)


def _extract(path: Path) -> tuple[str, str, int | None]:
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _extract_pdf(path)
        if suffix == ".docx":
            return _extract_docx(path)
        if suffix in {".txt", ".md"}:
            return _extract_text_file(path)
        if suffix == ".csv":
            return _extract_csv(path)
        if suffix in {".xlsx", ".xls"}:
            return _extract_xlsx(path)
        if suffix == ".json":
            return _extract_json(path)
        return "", "saved_only", None
    except Exception:
        return "", "read_error", None


def process_uploaded_files(uploaded_files: Iterable[Any], upload_dir: Path) -> list[dict[str, Any]]:
    files = list(uploaded_files or [])
    if len(files) > MAX_FILES:
        raise ValueError(f"Можно загрузить не более {MAX_FILES} файлов за один запрос.")

    sizes = [len(file.getvalue()) for file in files]
    if any(size > MAX_FILE_BYTES for size in sizes):
        raise ValueError("Размер одного файла не должен превышать 20 МБ.")
    if sum(sizes) > MAX_TOTAL_BYTES:
        raise ValueError("Общий размер файлов не должен превышать 60 МБ.")

    upload_dir.mkdir(parents=True, exist_ok=True)
    result: list[dict[str, Any]] = []

    for uploaded_file, size in zip(files, sizes):
        content = uploaded_file.getvalue()
        safe_name = _safe_name(uploaded_file.name)
        suffix = Path(safe_name).suffix
        stem = Path(safe_name).stem
        digest = hashlib.sha256(content).hexdigest()[:10]
        stored_name = f"{stem}-{digest}{suffix}"
        destination = upload_dir / stored_name
        destination.write_bytes(content)

        text, status, page_count = _extract(destination)
        clean_text = _clean_text(text)
        result.append(
            {
                "name": uploaded_file.name,
                "stored_name": stored_name,
                "local_path": str(destination),
                "mime_type": uploaded_file.type or mimetypes.guess_type(uploaded_file.name)[0] or "application/octet-stream",
                "size_bytes": size,
                "page_count": page_count,
                "status": status,
                "text": text[:MAX_TEXT_CHARS],
                "excerpt": clean_text[:MAX_EXCERPT_CHARS],
                "sha256": hashlib.sha256(content).hexdigest(),
            }
        )

    return result


def metadata_for_storage(documents: list[dict[str, Any]]) -> list[dict[str, Any]]:
    allowed = {
        "name", "stored_name", "local_path", "mime_type", "size_bytes",
        "page_count", "status", "excerpt", "sha256"
    }
    return [{key: value for key, value in document.items() if key in allowed} for document in documents]


def documents_for_backend(documents: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        {
            "name": document.get("name", ""),
            "mime_type": document.get("mime_type", ""),
            "size_bytes": document.get("size_bytes", 0),
            "page_count": document.get("page_count"),
            "status": document.get("status", ""),
            "text": document.get("text", "")[:MAX_TEXT_CHARS],
        }
        for document in documents
    ]
